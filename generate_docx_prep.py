import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_document():
    doc = Document()
    
    # Configure styles & margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base colors
    navy_color = RGBColor(12, 35, 64)       # Primary
    teal_color = RGBColor(0, 128, 128)     # Secondary
    gray_color = RGBColor(100, 100, 100)   # Muted text
    
    # Custom styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(51, 51, 51)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("VIDEO PLAYBACK TRACKER — REVERSE-ENGINEERED STUDY GUIDE")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = navy_color
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("A High-Performance Cross-Browser Extension & Neo-Brutalist Sync Dashboard\nCompiled by a Senior Software Engineer")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = gray_color
    
    doc.add_paragraph("").paragraph_format.space_after = Pt(20)
    
    # Helper to add headings
    def add_custom_heading(text, level, space_before=18, space_after=6):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(space_before)
        h.paragraph_format.space_after = Pt(space_after)
        h.paragraph_format.keep_with_next = True
        
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.bold = True
        
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = navy_color
            pPr = h._p.get_or_add_pPr()
            pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                             r'<w:bottom w:val="single" w:sz="12" w:space="4" w:color="0C2340"/>'
                             r'</w:pBdr>')
            pPr.append(pBdr)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = teal_color
        else:
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(12, 35, 64)
            
        return h

    # Section 1: PROJECT OVERVIEW
    add_custom_heading("1. PROJECT OVERVIEW", 1)
    
    doc.add_paragraph(
        "\"Rewind\" is a cross-browser extension for Google Chrome and Mozilla Firefox that seamlessly records "
        "and synchronizes video playback timestamps across multiple streaming platforms (Netflix, YouTube, Crunchyroll) "
        "into a centralized Neo-Brutalist user dashboard, enabling students to resume playback at their precise second."
    )
    
    doc.add_paragraph().add_run("Real-World Problem It Solves:").bold = True
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Fragmented Streaming Histories: ").bold = True
    p.add_run("Different media players do not communicate. If a user pauses a tutorial on YouTube and catches up on a documentary on Netflix, they must manually remember timestamps, scrubbing through timelines repeatedly.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Single-Page Application (SPA) Tracking: ").bold = True
    p.add_run("Modern streaming platforms do not trigger standard page reloads on video transitions, blinding typical DOM tracking extensions.")

    doc.add_paragraph().add_run("Key Features:").bold = True
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("MutationObserver SPA Capture: ").bold = True
    p.add_run("Monitors structural DOM modifications to track video mount lifecycles without reloading.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("WeakSet Memory Optimization: ").bold = True
    p.add_run("Locks video elements in weak references. Automatically frees memory during garbage collection to prevent leaks.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Sandboxed Neural Sync: ").bold = True
    p.add_run("Pairs sandboxed browser extensions with the dashboard with zero user login screens.")

    # Section 2: COMPLETE TECH STACK BREAKDOWN
    add_custom_heading("2. COMPLETE TECH STACK BREAKDOWN", 1)
    
    table = doc.add_table(rows=6, cols=5)
    table.style = 'Light Shading Accent 1'
    
    hdr = table.rows[0].cells
    hdr[0].text = 'Technology'
    hdr[1].text = 'Role in Project'
    hdr[2].text = 'Chosen For'
    hdr[3].text = 'Alternatives'
    hdr[4].text = 'Trade-offs'
    
    rows_data = [
        ('Browser APIs', 'MutationObserver & WeakSet.', 'Direct DOM observation, self-cleaning element tracking in SPAs, preventing memory leaks.', 'Interval polling', 'Precise hook callbacks vs. initial performance parsing.'),
        ('Chrome MV3', 'Service worker architecture.', 'Chrome Manifest V3 compliance, granular background script isolation.', 'Manifest V2', 'Long-term store support vs. ephemeral service worker state cycles.'),
        ('Firefox MV2', 'Content script execution.', 'Firefox compatibility, explicit Content Security Policy whitelists.', 'Manifest V3', 'Stable connection support vs. legacy syntax compilation.'),
        ('React + Vite', 'Neo-Brutalist Dashboard.', 'Spring list transitions, fast HMR updates, high-performance UI state manipulation.', 'Next.js', 'Instant client compilation speeds vs. lacking SSR routing optimization.'),
        ('Firestore', 'Dual-bucket sync pipeline.', 'Real-time WebSocket snapshot listeners, temporary buffer bucket processing.', 'PostgreSQL, REST APIs', 'WebSocket updates vs. client structural validation rules.')
    ]
    
    for i, data in enumerate(rows_data, start=1):
        row = table.rows[i].cells
        row[0].text = data[0]
        row[1].text = data[1]
        row[2].text = data[2]
        row[3].text = data[3]
        row[4].text = data[4]

    doc.add_paragraph("").paragraph_format.space_after = Pt(10)

    # Section 3: FULL ARCHITECTURE WALKTHROUGH
    add_custom_heading("3. FULL ARCHITECTURE WALKTHROUGH", 1)
    
    p = doc.add_paragraph()
    p.add_run("Step-by-Step Timestamp Sync Lifecycle:\n").bold = True
    p.add_run("1. DOM Capture: ").bold = True
    p.add_run("MutationObserver spots a video tag on Netflix. It registers listeners for `play`, `pause`, and `ended` events, and adds the DOM reference to a WeakSet.\n")
    p.add_run("2. Payload Buffering: ").bold = True
    p.add_run("When the user pauses or triggers `beforeunload` (tab close), the extension pushes a JSON timestamp payload to the temporary Firestore buffer bucket.\n")
    p.add_run("3. Gatekeeping: ").bold = True
    p.add_run("The fully authenticated web dashboard reads the buffer snapshot, replaces client-side parameters with server atomic timestamps, writes the verified item to permanent history, and purges the buffer.")

    # Section 4: DATABASE DESIGN
    add_custom_heading("4. DATABASE DESIGN", 1)
    
    db_table = doc.add_table(rows=3, cols=3)
    db_table.style = 'Light Shading Accent 1'
    
    hdr = db_table.rows[0].cells
    hdr[0].text = 'Collection'
    hdr[1].text = 'Primary Fields'
    hdr[2].text = 'Relationships & Rationale'
    
    schemas = [
        ('sync_buffer', 'bufferId: String (PK)\nuserId: String\nvideoUrl: String\ntimestamp: Number\ntitle: String\nplatform: String', 'Temporary Dropbox. Extension writes loosely; dashboard acts as gatekeeper.'),
        ('watch_history', 'historyId: String (PK)\nuserId: String\nvideoUrl: String\ntimestamp: Number\ntitle: String\nserverTime: ServerTimestamp', 'Permanent watch database. Highly structured, validated, sorted descending.')
    ]
    
    for i, data in enumerate(schemas, start=1):
        row = db_table.rows[i].cells
        row[0].text = data[0]
        row[1].text = data[1]
        row[2].text = data[2]

    doc.add_paragraph("").paragraph_format.space_after = Pt(10)

    # Section 5: AUTHENTICATION + SECURITY
    add_custom_heading("5. AUTHENTICATION + SECURITY", 1)
    doc.add_paragraph(
        "Bypassing Cookie Sandboxes via Neural Sync:\n"
        "Extensions cannot read HttpOnly dashboard cookies. On login, the dashboard embeds the decrypted "
        "user ID in a hidden DOM element. The extension's content script scrapes this element on the dashboard domain, "
        "relays the ID to the background worker via `chrome.runtime.sendMessage`, and establishes pairing securely."
    )

    # Section 6: CORE CONCEPTS I MUST UNDERSTAND
    add_custom_heading("6. CORE CONCEPTS I MUST UNDERSTAND", 1)
    p = doc.add_paragraph()
    p.add_run("1. JavaScript WeakSet: ").bold = True
    p.add_run("Holds object references weakly. If the browser destroys a video tag on Netflix, the garbage collector "
              "removes it from our WeakSet, keeping our memory leak-free.")
    p = doc.add_paragraph()
    p.add_run("2. Content Security Policy (CSP): ").bold = True
    p.add_run("Firefox requires explicit `connect-src` manifests to allow WebSockets and API connections, "
              "preventing silent connection failures.")

    # Section 7: MOST IMPORTANT CODE SECTIONS
    add_custom_heading("7. MOST IMPORTANT CODE SECTIONS", 1)
    doc.add_paragraph(
        "content.js: Initializes MutationObserver, tracks video components, and parses metadata selectors. "
        "Understand metadata selector loops and why WeakSet is used."
    )

    # Section 8: HARDEST TECHNICAL CHALLENGES
    add_custom_heading("8. HARDEST TECHNICAL CHALLENGES", 1)
    doc.add_paragraph(
        "Challenge: Solving the Sandboxed Tab-Closing Event Capture\n"
        "The Problem: Closing a tab kills JavaScript before `pause` runs. Traditional beforeunload calls are often "
        "blocked by modern browser sandboxes.\n"
        "The Solution: We built a dual redundancy pipeline: a 30-second low-impact periodic sync loop to localStorage, "
        "combined with a standard `beforeunload` callback that evaluates play state and fires off a last-second, "
        "optimized fetch API request before thread destruction."
    )

    # Section 9: INTERVIEW PREPARATION
    add_custom_heading("9. INTERVIEW PREPARATION", 1)
    p = doc.add_paragraph()
    p.add_run("Q: What is the 'Ghost Deletion Bug' and how did you debug and resolve it?\n").bold = True
    p.add_run("Answer: ").italic = True
    p.add_run("The extension generated temporary client IDs. The dashboard spread fetched documents: `{ id: doc.id, ...doc.data() }`. "
              "Because JS spreads right-to-left, the client's temporary ID overwrote the database ID. The delete "
              "call tried to delete non-existent IDs. I resolved it by swapping the order: `{ ...doc.data(), id: doc.id }`.")

    # Section 10: QUICK REVISION NOTES
    add_custom_heading("10. QUICK REVISION NOTES", 1)
    doc.add_paragraph().add_run("30-Second Pitch:").bold = True
    doc.add_paragraph(
        "\"Rewind is a cross-browser extension for Chrome and Firefox that automatically saves video playback timestamps "
        "across streaming platforms. Using MutationObservers and WeakSets, it captures states without memory leaks in SPAs, "
        "syncing data to a Neo-Brutalist React dashboard via a secure, sandboxed dual-bucket Firestore sync pipeline.\""
    )

    # Section 11: RED FLAGS / THINGS THAT MAY EXPOSE ME
    add_custom_heading("11. RED FLAGS / THINGS THAT MAY EXPOSE ME", 1)
    doc.add_paragraph(
        "Red Flag: Saying extensions can freely access data from other tabs.\n"
        "Fix: Explain that extensions run inside sandbox environments, necessitating explicit permission manifests "
        "and message brokers."
    )

    # Section 12: LEARNING PRIORITY
    add_custom_heading("12. LEARNING PRIORITY", 1)
    doc.add_paragraph("1. Highest: MutationObserver and DOM tree parsing.")
    doc.add_paragraph("2. Medium: Content Security Policies and Extension message routing.")
    doc.add_paragraph("3. Low: Neo-Brutalist UI styles.")

    doc.save("/Users/aadeshkhande/Documents/Professional/Own/Video PlayBack Tracker/Video_Playback_Tracker_Interview_Prep.docx")

if __name__ == '__main__':
    create_document()
    print("Video_Playback_Tracker_Interview_Prep.docx has been created successfully!")
